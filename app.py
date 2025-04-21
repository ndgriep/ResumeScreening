import google.generativeai as genai
import PyPDF2 as pdf
import docx as doc
import json
import os
from dotenv import load_dotenv
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, ttk
import threading
import re
import datetime
import csv

#  Configuration and Setup 
load_dotenv()

# Global variables
selected_resume_path = None
analysis_history = []  # Store previous analyses
current_resume_data = None  # Store extracted resume data for reuse

def initialize_api():
    """Initialize the Gemini API with proper error handling."""
    try:
        # Check environment variable name for the API key
        gemini_api_key = os.getenv("GOOGLE_API_KEY")
        if not gemini_api_key:
            raise ValueError("API key not found in environment variables. Please set GOOGLE_API_KEY in your .env file.")
        
        genai.configure(api_key=gemini_api_key)
        return genai.GenerativeModel('gemini-2.0-flash')
    except Exception as e:
        print(f"Error configuring Gemini API: {e}")
        messagebox.showerror("API Error", f"Failed to configure Gemini API: {e}\nPlease ensure your .env file has a valid API key.")
        return None

# --- Text Extraction Functions ---
def extract_text_from_pdf(pdf_path):
    """Extract text from PDF files with robust error handling."""
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = pdf.PdfReader(file)
            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception as decrypt_error:
                    print(f"Warning: Could not decrypt PDF '{os.path.basename(pdf_path)}': {decrypt_error}")
                    return f"Error: Could not decrypt PDF '{os.path.basename(pdf_path)}'."
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_error:
                    print(f"Warning: Error extracting text from page {page_num} in PDF: {page_error}")
                    
            if not text.strip():
                print("Warning: No text extracted from PDF. The file might be scanned or image-based.")
                return "Error: Could not extract text from PDF. The file might be scanned or image-based."
                
    except FileNotFoundError:
        print(f"Error: PDF file not found at {pdf_path}")
        return None
    except Exception as e:
        print(f"Error extracting text from PDF '{os.path.basename(pdf_path)}': {e}")
        if "EOF marker not found" in str(e):
            print("Suggestion: The PDF might be corrupted or incomplete.")
        return f"Error extracting text from PDF: {e}"
    return text

def extract_text_from_docx(docx_path):
    """Extract text from DOCX files with robust error handling."""
    text = ""
    try:
        document = doc.Document(docx_path)
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        if not text.strip():
            print("Warning: No text extracted from DOCX. The file might be empty or contain only images.")
            return "Error: No text content found in the DOCX file."
            
    except FileNotFoundError:
        print(f"Error: DOCX file not found at {docx_path}")
        return None
    except Exception as e:
        print(f"Error extracting text from DOCX '{os.path.basename(docx_path)}': {e}")
        return f"Error extracting text from DOCX: {e}"
    return text

def extract_text_from_txt(txt_path):
    """Extract text from TXT files with robust error handling and multiple encoding support."""
    text = ""
    try:
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'utf-16', 'ascii']
        for encoding in encodings_to_try:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    text = file.read()
                print(f"Successfully read file with encoding: {encoding}")
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError(f"Could not decode file with tried encodings: {encodings_to_try}")
            
        if not text.strip():
            print("Warning: File is empty.")
            return "Error: The text file is empty."
            
    except FileNotFoundError:
        print(f"Error: TXT file not found at {txt_path}")
        return None
    except Exception as e:
        print(f"Error extracting text from TXT '{os.path.basename(txt_path)}': {e}")
        return f"Error extracting text from TXT: {e}"
    return text

# --- Gemini API Interaction Functions ---
def clean_json_string(json_string):
    """Clean JSON string from API responses to make it parseable."""
    # Remove markdown code blocks if present
    cleaned = re.sub(r'```(?:json)?\s*|\s*```', '', json_string.strip())
    # Handle potential trailing commas in JSON objects/arrays (common API mistake)
    cleaned = re.sub(r',\s*}', '}', cleaned)
    cleaned = re.sub(r',\s*]', ']', cleaned)
    return cleaned

def extract_resume_info(resume_text, model):
    """Extract structured information from resume text using Gemini API."""
    if not resume_text or resume_text.strip().startswith("Error:"):
        print("Skipping resume info extraction due to empty or error text.")
        return None
        
    prompt = f"""Analyze the following resume text and extract the specified information into a structured JSON format.
    Resume text:
    {resume_text}
    
    Please extract and return the following information as valid JSON:
    {{
      "personal_info": {{
        "name": "Full name of the candidate",
        "email": "Email address if available",
        "phone": "Phone number if available",
        "location": "Location/address if available"
      }},
      "education": [
        {{
          "degree": "Degree name",
          "institution": "Institution name",
          "dates": "Date range",
          "gpa": "GPA if available"
        }}
      ],
      "work_experience": [
        {{
          "title": "Job title",
          "company": "Company name",
          "dates": "Employment date range",
          "responsibilities": ["Key responsibility 1", "Key responsibility 2", "..."]
        }}
      ],
      "skills": ["Skill 1", "Skill 2", "..."],
      "certifications": ["Certification 1", "Certification 2", "..."],
      "summary": "Brief professional summary extracted from the resume"
    }}
    
    Format the output as valid JSON with no additional text or explanation. Handle missing information by using empty strings or arrays as appropriate, but maintain the structure.
    """
    
    try:
        response = model.generate_content(prompt)
        json_string = clean_json_string(response.text)
        
        try:
            extracted_data = json.loads(json_string)
            return extracted_data
        except json.JSONDecodeError as json_e:
            print(f"Error decoding JSON response: {json_e}")
            # Attempt to fix common JSON issues
            print("Attempting to fix JSON format...")
            sanitized_json = re.sub(r'(\w+)(?=\s*:)', r'"\1"', json_string)  # Quote unquoted keys
            try:
                extracted_data = json.loads(sanitized_json)
                return extracted_data
            except:
                return {"error": "Failed to parse resume data from API response.", "raw_response": response.text}
                
    except Exception as e:
        print(f"Error calling Gemini API during resume extraction: {e}")
        return {"error": f"API call failed during resume extraction: {e}"}

def analyze_resume_against_job(resume_data, job_description, model):
    """Analyze how well the resume matches the job description using Gemini API."""
    if not resume_data or "error" in resume_data:
        print("Skipping analysis because resume data is missing or contains errors.")
        return {"error": "Cannot perform analysis due to issues with resume data."}
        
    if not job_description or not job_description.strip():
        print("Skipping analysis because job description is empty.")
        return {"error": "Job description is empty."}
        
    resume_json_string = json.dumps(resume_data, indent=2)
    
    prompt = f"""Analyze the provided resume (in JSON format) against the following job description.
    
    RESUME:
    {resume_json_string}
    
    JOB DESCRIPTION:
    {job_description}
    
    Provide a detailed analysis in JSON format with the following structure:
    {{
      "match_score": "A score from 0 to 100 indicating how well the resume matches the job description",
      "strengths": ["Strength 1", "Strength 2", "..."],
      "weaknesses": ["Weakness 1", "Weakness 2", "..."],
      "missing_skills": ["Missing skill 1", "Missing skill 2", "..."],
      "matching_skills": ["Matching skill 1", "Matching skill 2", "..."],
      "recommendations": ["Recommendation 1", "Recommendation 2", "..."],
      "summary": "A brief summary of the match analysis"
    }}
    
    Be thorough in your analysis, looking not just at exact keyword matches but also at relevant experience, transferable skills, and overall fit. Your recommendations should be specific and actionable for improving the resume to better match this specific job description.
    
    Format the output as valid JSON with no additional text or explanation.
    """
    
    try:
        response = model.generate_content(prompt)
        json_string = clean_json_string(response.text)
        
        try:
            analysis_data = json.loads(json_string)
            return analysis_data
        except json.JSONDecodeError as json_e:
            print(f"Error decoding JSON response during analysis: {json_e}")
            # Try to fix common JSON issues
            sanitized_json = re.sub(r'(\w+)(?=\s*:)', r'"\1"', json_string)  # Quote unquoted keys
            try:
                analysis_data = json.loads(sanitized_json)
                return analysis_data
            except:
                return {"error": "Failed to parse analysis data from API response.", "raw_response": response.text}
                
    except Exception as e:
        print(f"Error calling Gemini API during analysis: {e}")
        return {"error": f"API call failed during analysis: {e}"}

# --- File Processing ---
def extract_resume_text_from_file(file_path):
    """Extract text from a resume file based on its file extension."""
    try:
        if not file_path or not os.path.exists(file_path):
            print(f"Error: File not found or path is invalid: {file_path}")
            return None
            
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        else:
            print(f"Error: Unsupported file type: {file_extension}")
            messagebox.showwarning("Unsupported File", f"File type '{file_extension}' is not supported. Please use .pdf, .docx, or .txt.")
            return None
    except Exception as e:
        print(f"Error processing file '{os.path.basename(file_path)}': {e}")
        messagebox.showerror("File Processing Error", f"Could not process the file:\n{e}")
        return None

# --- Data Management Functions ---
def save_analysis_to_csv(analysis_data, resume_path, job_description):
    """Save analysis results to a CSV file for record keeping."""
    try:
        # Create a directory for saved analyses if it doesn't exist
        save_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "saved_analyses")
        os.makedirs(save_dir, exist_ok=True)
        
        # Create a timestamped filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        resume_filename = os.path.basename(resume_path)
        csv_filename = f"analysis_{timestamp}_{resume_filename}.csv"
        csv_path = os.path.join(save_dir, csv_filename)
        
        # Flatten the analysis data for CSV storage
        flattened_data = {
            "timestamp": timestamp,
            "resume_filename": resume_filename,
            "match_score": analysis_data.get("match_score", "N/A"),
            "summary": analysis_data.get("summary", ""),
            "job_description": job_description,
            "strengths": "; ".join(analysis_data.get("strengths", [])),
            "weaknesses": "; ".join(analysis_data.get("weaknesses", [])),
            "missing_skills": "; ".join(analysis_data.get("missing_skills", [])),
            "matching_skills": "; ".join(analysis_data.get("matching_skills", [])),
            "recommendations": "; ".join(analysis_data.get("recommendations", []))
        }
        
        # Write to CSV
        with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = flattened_data.keys()
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerow(flattened_data)
            
        print(f"Analysis saved to {csv_path}")
        return csv_path
    except Exception as e:
        print(f"Error saving analysis to CSV: {e}")
        return None

# --- GUI Functions ---
def select_file():
    """Open file dialog to select a resume file."""
    global selected_resume_path, current_resume_data
    file_path = filedialog.askopenfilename(
        title="Select Resume File", 
        filetypes=(("PDF files", "*.pdf"), ("Word files", "*.docx"), ("Text files", "*.txt"), ("All files", "*.*"))
    )
    
    if file_path:
        selected_resume_path = file_path
        file_label.config(text=f"Selected: {os.path.basename(file_path)}")
        # Reset resume data when new file is selected
        current_resume_data = None
        # Clear results
        clear_results()
    else:
        selected_resume_path = None
        file_label.config(text="No file selected.")

def clear_results():
    """Clear the results text area."""
    results_text.config(state=tk.NORMAL)
    results_text.delete('1.0', tk.END)
    results_text.config(state=tk.DISABLED)

def run_analysis():
    """Main function to run the resume analysis."""
    global selected_resume_path, current_resume_data
    
    if not selected_resume_path:
        messagebox.showwarning("Input Missing", "Please select a resume file first.")
        return
        
    job_desc = job_desc_text.get("1.0", tk.END).strip()
    if not job_desc:
        messagebox.showwarning("Input Missing", "Please paste the job description.")
        return
    
    # Initialize or get the model
    model = initialize_api()
    if not model:
        return
    
    # Update UI to show processing
    analyze_button.config(state=tk.DISABLED, text="Analyzing...")
    progress_bar.pack(pady=5)
    progress_bar.start(10)
    clear_results()
    update_results("Processing...\n\n")
    
    # Start analysis in a separate thread to keep UI responsive
    analysis_thread = threading.Thread(
        target=perform_analysis_thread, 
        args=(selected_resume_path, job_desc, model),
        daemon=True
    )
    analysis_thread.start()

def perform_analysis_thread(resume_path, job_description, model):
    """Run the analysis process in a separate thread."""
    global current_resume_data
    
    try:
        # Step 1: Extract text from resume
        update_results("Step 1: Extracting text from resume...\n")
        resume_text = extract_resume_text_from_file(resume_path)
        
        if not resume_text:
            update_results("Error: Could not extract text from resume file. See console for details.\nAnalysis aborted.")
            reset_gui_state()
            return
            
        if resume_text.strip().startswith("Error:"):
            update_results(f"Error during text extraction:\n{resume_text}\nAnalysis aborted.")
            reset_gui_state()
            return
            
        update_results("Text extracted successfully.\n")

        # Step 2: Extract structured resume information
        update_results("\nStep 2: Extracting structured information from resume via API...\n")
        
        # If we already have resume data for this file, skip API call
        if current_resume_data is None:
            current_resume_data = extract_resume_info(resume_text, model)
            
        if not current_resume_data or "error" in current_resume_data:
            error_msg = current_resume_data.get('error', 'Unknown error during resume extraction.') if current_resume_data else "Failed to extract resume data."
            raw_response = current_resume_data.get('raw_response', '') if current_resume_data else ""
            update_results(f"Error during resume information extraction:\n{error_msg}\n{raw_response}\nAnalysis aborted.")
            reset_gui_state()
            return
            
        update_results("Resume information extracted successfully.\n")

        # Step 3: Analyze resume against job description
        update_results("\nStep 3: Analyzing resume against job description via API...\n")
        analysis_results = analyze_resume_against_job(current_resume_data, job_description, model)
        
        if not analysis_results or "error" in analysis_results:
            error_msg = analysis_results.get('error', 'Unknown error during analysis.') if analysis_results else "Analysis failed."
            raw_response = analysis_results.get('raw_response', '') if analysis_results else ""
            update_results(f"Error during job match analysis:\n{error_msg}\n{raw_response}\nAnalysis finished with errors.")
            reset_gui_state()
            return
        
        # Format and display results
        update_results("\n--- Analysis Results ---\n")
        update_results(f"Match Score: {analysis_results.get('match_score', 'N/A')}/100\n\n")
        
        if "strengths" in analysis_results and analysis_results["strengths"]:
            update_results("Strengths:\n")
            for strength in analysis_results["strengths"]:
                update_results(f"✓ {strength}\n")
            update_results("\n")
        
        if "weaknesses" in analysis_results and analysis_results["weaknesses"]:
            update_results("Areas for Improvement:\n")
            for weakness in analysis_results["weaknesses"]:
                update_results(f"! {weakness}\n")
            update_results("\n")
        
        if "missing_skills" in analysis_results and analysis_results["missing_skills"]:
            update_results("Missing Skills:\n")
            for skill in analysis_results["missing_skills"]:
                update_results(f"- {skill}\n")
            update_results("\n")
        
        update_results("Recommendations:\n")
        for rec in analysis_results.get('recommendations', ['No recommendations provided.']):
            update_results(f"• {rec}\n")
        
        if "summary" in analysis_results:
            update_results(f"\nSummary:\n{analysis_results['summary']}\n")
        
        update_results("\n--- Analysis Complete ---")
        
        # Save analysis to history
        analysis_entry = {
            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "resume_path": resume_path,
            "resume_name": os.path.basename(resume_path),
            "job_description": job_description,
            "analysis_results": analysis_results,
            "resume_data": current_resume_data
        }
        analysis_history.append(analysis_entry)
        
        # Enable the save button
        root.after(0, lambda: save_button.config(state=tk.NORMAL))
        
    except Exception as e:
        print(f"An unexpected error occurred during analysis thread: {e}")
        update_results(f"\n--- An Unexpected Error Occurred ---\n{e}\nSee console for more details.")
        import traceback
        traceback.print_exc()
    finally:
        reset_gui_state()

def save_analysis_action():
    """Save the current analysis results to a CSV file."""
    global analysis_history, selected_resume_path
    
    if not analysis_history:
        messagebox.showwarning("No Analysis", "Please complete an analysis first.")
        return
    
    # Use the most recent analysis
    latest_analysis = analysis_history[-1]
    job_desc = job_desc_text.get("1.0", tk.END).strip()
    
    # Save to CSV
    csv_path = save_analysis_to_csv(latest_analysis["analysis_results"], selected_resume_path, job_desc)
    
    if csv_path:
        messagebox.showinfo("Analysis Saved", f"Analysis results saved to:\n{csv_path}")
    else:
        messagebox.showerror("Save Error", "Failed to save analysis results.")

def update_results(text):
    """Update the results text area safely using the main thread."""
    root.after(0, _do_update_results, text)

def _do_update_results(text):
    """Helper function to update results text from the main thread."""
    try:
        results_text.config(state=tk.NORMAL)
        results_text.insert(tk.END, text)
        results_text.see(tk.END)
        results_text.config(state=tk.DISABLED)
        root.update_idletasks()
    except tk.TclError as e:
        print(f"Tkinter error during update: {e}")

def reset_gui_state():
    """Reset GUI elements to their default state."""
    root.after(0, _do_reset_gui_state)

def _do_reset_gui_state():
    """Helper function to reset GUI from the main thread."""
    try:
        analyze_button.config(state=tk.NORMAL, text="Analyze Resume vs. Job Description")
        progress_bar.stop()
        progress_bar.pack_forget()
    except tk.TclError as e:
        print(f"Tkinter error during reset: {e}")

def show_about_dialog():
    """Show information about the application."""
    about_text = """Resume Analyzer v2.0

This application helps you analyze your resume against job descriptions using AI.

Features:
- Extract text from PDF, DOCX, and TXT resumes
- Analyze resume against job descriptions
- Save analysis results for future reference

Powered by Google's Gemini API
"""
    messagebox.showinfo("About Resume Analyzer", about_text)

def show_help_dialog():
    """Show help information for using the application."""
    help_text = """How to Use Resume Analyzer:

1. Select your resume file (.pdf, .docx, or .txt)
2. Paste the job description text
3. Click "Analyze" to compare your resume to the job
4. Review the analysis results
5. Save the analysis for your records

Tips:
- For best results, use a well-formatted resume
- Make sure job descriptions contain detailed requirements
"""
    messagebox.showinfo("Help", help_text)

def view_previous_analyses():
    """Open a window to view previous analysis history."""
    if not analysis_history:
        messagebox.showinfo("No History", "No analyses have been performed yet.")
        return
    
    # Create a new window
    history_window = tk.Toplevel(root)
    history_window.title("Analysis History")
    history_window.geometry("600x500")
    
    # Create a frame for the list
    list_frame = tk.Frame(history_window)
    list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    # Create a label
    tk.Label(list_frame, text="Previous Analyses:").pack(anchor="w")
    
    # Create a listbox with scrollbar
    listbox_frame = tk.Frame(list_frame)
    listbox_frame.pack(fill=tk.BOTH, expand=True)
    
    scrollbar = tk.Scrollbar(listbox_frame)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    history_listbox = tk.Listbox(listbox_frame, yscrollcommand=scrollbar.set)
    history_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    scrollbar.config(command=history_listbox.yview)
    
    # Populate the listbox
    for i, entry in enumerate(analysis_history):
        display_text = f"{entry['timestamp']} - {entry['resume_name']} (Score: {entry['analysis_results'].get('match_score', 'N/A')})"
        history_listbox.insert(tk.END, display_text)
    
    # Create detail view
    detail_frame = tk.Frame(history_window)
    detail_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
    
    tk.Label(detail_frame, text="Details:").pack(anchor="w")
    
    detail_text = scrolledtext.ScrolledText(detail_frame, wrap=tk.WORD, height=10)
    detail_text.pack(fill=tk.BOTH, expand=True)
    
    # Function to display selected entry details
    def on_select(event):
        if history_listbox.curselection():
            index = history_listbox.curselection()[0]
            entry = analysis_history[index]
            
            detail_text.delete('1.0', tk.END)
            detail_text.insert(tk.END, f"Date/Time: {entry['timestamp']}\n")
            detail_text.insert(tk.END, f"Resume: {entry['resume_name']}\n")
            detail_text.insert(tk.END, f"Match Score: {entry['analysis_results'].get('match_score', 'N/A')}\n\n")
            
            detail_text.insert(tk.END, "Summary:\n")
            detail_text.insert(tk.END, f"{entry['analysis_results'].get('summary', 'No summary available.')}\n\n")
            
            detail_text.insert(tk.END, "Key Strengths:\n")
            for strength in entry['analysis_results'].get('strengths', ['No strengths listed.']):
                detail_text.insert(tk.END, f"- {strength}\n")
    
    history_listbox.bind('<<ListboxSelect>>', on_select)
    
    # Select the first item by default
    if analysis_history:
        history_listbox.select_set(0)
        history_listbox.event_generate('<<ListboxSelect>>')

# --- GUI Setup ---
root = tk.Tk()
root.title("Resume Analyzer")
root.geometry("900x750")

# Add a menu bar
menu_bar = tk.Menu(root)
root.config(menu=menu_bar)

# File menu
file_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="Select Resume", command=select_file)
file_menu.add_command(label="Save Analysis", command=save_analysis_action)
file_menu.add_command(label="View Previous Analyses", command=view_previous_analyses)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=root.quit)

# Help menu
help_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Help", menu=help_menu)
help_menu.add_command(label="How to Use", command=show_help_dialog)
help_menu.add_command(label="About", command=show_about_dialog)

# Main frame
main_frame = tk.Frame(root, padx=10, pady=10)
main_frame.pack(fill=tk.BOTH, expand=True)

# File selection frame
file_frame = tk.LabelFrame(main_frame, text="Resume File", padx=5, pady=5)
file_frame.pack(fill=tk.X, pady=(0, 10))

select_button = tk.Button(file_frame, text="Select Resume File (.pdf, .docx, .txt)", command=select_file)
select_button.pack(side=tk.LEFT, padx=10, pady=5)

file_label = tk.Label(file_frame, text="No file selected.", anchor="w")
file_label.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)

# Job description frame
job_desc_frame = tk.LabelFrame(main_frame, text="Job Description", padx=5, pady=5)
job_desc_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

job_desc_text = scrolledtext.ScrolledText(job_desc_frame, wrap=tk.WORD, height=8, borderwidth=1, relief="solid")
job_desc_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

# Analysis button frame
buttons_frame = tk.Frame(main_frame)
buttons_frame.pack(fill=tk.X, pady=(0, 10))

analyze_button = tk.Button(
    buttons_frame, 
    text="Analyze Resume vs. Job Description", 
    command=run_analysis,
    bg="#4CAF50",
    fg="white",
    height=2
)
analyze_button.pack(fill=tk.X, expand=True)

# Progress bar (hidden by default)
progress_bar = ttk.Progressbar(main_frame, mode='indeterminate')
# Will be packed when needed

# Results frame
results_frame = tk.LabelFrame(main_frame, text="Results", padx=5, pady=5)
results_frame.pack(fill=tk.BOTH, expand=True)

results_text = scrolledtext.ScrolledText(
    results_frame, 
    wrap=tk.WORD, 
    height=15, 
    state=tk.DISABLED, 
    borderwidth=1, 
    relief="solid",
    font=("Segoe UI", 10)
)
results_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

# Save button at the bottom
save_button = tk.Button(
    main_frame,
    text="Save Analysis Results",
    command=save_analysis_action,
    state=tk.DISABLED
)
save_button.pack(pady=10)

# Run the main loop
root.mainloop()